import streamlit as st
import os
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from localApp import generate_text_from_model
from extract_high_freq_words import extract_high_freq_words_from_file
from translate import Translator

translator = Translator(from_lang='zh-cn', to_lang='en')
doc = Document()


# 添加word一级标题
def add_MainHeading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


# 添加word正文内容
def add_body(text):
    content = (text)
    p = doc.add_paragraph(content)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


# 创建大纲word文档
def create_file(course_name, course_alltime, course_labtime, course_sub, course_type, file):
    st.warning('Submission successful, course outline is being generated, please wait...')

    # 创建一个新的Word文档

    st.warning("The 'Course Objectives' section of the course outline is being generated...")

    # 添加word标题，并设置字体为宋体
    # title = doc.add_heading(course_name + '课程' + '教学大纲', level=0)
    title = doc.add_heading("The course outline of " + course_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.bold = True  # 标题加粗

    prompt1 = "将" + course_name + "翻译成英语，只输出英文结果即可，不要其他任何符号和描述"
    # course_nameE = generate_text_from_model(prompt1)

    # 定义信息字段和对应的示例数据
    course_info = {
        '课程名称': course_name,
        # '英文名称': course_nameE,
        '学时': course_alltime,
        '实验学时': course_labtime,
        '课程性质': course_type,
        '适用专业': course_sub,
    }

    # 添加一个表格，表格的行数是字段数，列数是2
    table = doc.add_table(rows=len(course_info), cols=2)

    # 设置表格样式（可选）
    table.style = 'Table Grid'

    # 填充信息字段和数据
    for i, (field, value) in enumerate(course_info.items()):
        row_cells = table.rows[i].cells
        for idx, text in enumerate([field, value]):
            paragraph = row_cells[idx].paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = text
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            if idx == 0:  # 字段名加粗
                run.bold = True
            run.font.size = Pt(12)

    st.success("The 'Course Instructions' section of the course outline has been generated successfully!")

    high_freq_words = extract_high_freq_words_from_file(file, 20)

    # 添加一个空行
    doc.add_paragraph()

    st.warning("The 'Course Instructions' section of the course outline is being generated...")

    # 添加课程说明标题
    add_MainHeading('Ⅰ Course description')

    prompt3 = "输出" + course_name + "这一门大学课程的课程说明，80字以内，课程类型为" + course_type + ",适用专业为" + course_sub
    course_anal = generate_text_from_model(prompt3)
    course_anal = translator.translate(course_anal)

    # 添加课程说明内容
    add_body(course_anal)

    st.success("The 'Course Instructions' section of the course outline has been generated successfully!")

    # 添加一个空行
    doc.add_paragraph()

    st.warning("The 'Course Objectives' section of the course outline is being generated...")

    # 添加课程目标标题
    add_MainHeading('Ⅱ Course objectives')

    prompt4 = "输出" + course_name + "这一门大学课程的课程目标，分为四个点，按照下面的格式，目标1：了解什么内容，目标2：理解什么内容，目标3：掌握什么内容，目标4：运用什么内容，只分四段输出四个目标，不要任何符号和描述"
    course_tar = generate_text_from_model(prompt4)
    course_tar = translator.translate(course_tar)

    # 添加课程目标内容
    add_body(course_tar)

    st.success("The 'Course Objectives' section of the course outline has been generated successfully!")

    # 添加一个空行
    doc.add_paragraph()

    st.warning("The 'Course Content and Schedule' section of the course outline is being generated...")

    # 添加教学内容标题
    add_MainHeading('Ⅲ Teaching content and class schedule')

    prompt5 = "设计" + course_name + "这一门大学课程的教学内容与学时安排，课程类型为" + course_type + ",适用专业为" + course_sub + "，200字以内"

    course_txt = generate_text_from_model(prompt5)
    course_txt = translator.translate(course_txt)

    # 添加教学内容
    add_body(course_txt)

    st.success("The 'Course Content and Schedule' section of the course outline has been generated successfully!")

    # 添加一个空行
    doc.add_paragraph()

    st.warning("The 'Teaching Methods' section of the course outline is being generated...")

    # 添加教学方法标题
    add_MainHeading('Ⅳ Teaching method')

    prompt6 = "设计" + course_name + "这一门大学课程的教学方式，不要任何符号和描述，只输出答案，80字以内"
    course_teachway = generate_text_from_model(prompt6)
    course_teachway = translator.translate(course_teachway)

    # 添加教学方法内容
    add_body(course_teachway)

    st.success("The 'Teaching Methods' section of the course outline has been generated successfully!")

    # 添加一个空行
    doc.add_paragraph()

    st.warning("The 'Assessment Methods' section of the course outline is being generated...")

    # 添加考核方式标题
    add_MainHeading('Ⅴ Assessment method')

    prompt7 = "设计" + course_name + "这一门大学课程的考核方式，不要任何符号和描述，只输出答案，80字以内。"
    course_exam = generate_text_from_model(prompt7)
    course_exam = translator.translate(course_exam)

    # 添加考核方式内容
    add_body(
        "This course aims to cultivate students' theoretical literacy, practical ability and innovative thinking. " + course_exam)

    st.success("The 'Assessment Methods' section of the course outline has been generated successfully!")

    # 保存文档
    # doc.save('result\\' + course_name + ' course outline.docx')
    st.success('The course outline has been generated successfully!')

    # doc.save( course_name + '教学大纲.docx')

    time.sleep(1)

    # os.startfile('result')
    # os.startfile('result\\' + course_name + ' course outline.docx')
    os.startfile(
        r"C:\Users\Intel\Desktop\demo论文\5.9\class_assistant\courseware\result\Introduction to Information System course outline.docx")
